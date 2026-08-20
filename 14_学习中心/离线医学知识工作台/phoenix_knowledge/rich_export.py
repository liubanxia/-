from __future__ import annotations

import html
import os
import re
import shutil
from dataclasses import dataclass
from pathlib import Path


_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\((.+)\)")
_LINK_RE = re.compile(r"\[([^\]]+)\]\((.+)\)")
_MARKUP_RE = re.compile(r"[*_`~]+")
_SAFE_RE = re.compile(r'[\\/:*?"<>|\r\n]+')
_PDF_PAGE_CHAR_LIMIT = 1000
_PDF_PAGE_LINE_LIMIT = 28
_PDF_LINE_CHAR_LIMIT = 900


@dataclass(frozen=True)
class ExportBundle:
    output_dir: Path
    markdown: Path
    text: Path
    docx: Path
    pdf: Path

    @property
    def output_paths(self) -> tuple[Path, ...]:
        return (self.pdf, self.docx, self.markdown, self.text)


def _safe_name(value: str) -> str:
    return (_SAFE_RE.sub("_", value).strip(" ._") or "Phoenix医学专题")[:96]


def _resolve_image(base_dir: Path, reference: str) -> Path | None:
    raw = (reference or "").strip().strip("<>")
    if not raw or raw.startswith(("http://", "https://", "data:")):
        return None
    candidate = Path(raw)
    if not candidate.is_absolute():
        candidate = base_dir / candidate
    try:
        candidate = candidate.resolve()
    except OSError:
        pass
    return candidate if candidate.is_file() else None


def markdown_to_plain(markdown: str) -> str:
    text = (markdown or "").replace("\r\n", "\n")
    text = _IMAGE_RE.sub(
        lambda m: f"[图像：{m.group(1) or Path(m.group(2)).name}]",
        text,
    )
    text = _LINK_RE.sub(lambda m: m.group(1), text)
    lines: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if line.startswith("```"):
            continue
        line = re.sub(r"^\s{0,3}#{1,6}\s*", "", line)
        line = re.sub(r"^\s*>\s?", "", line)
        line = re.sub(r"^\s*[-*+]\s+", "• ", line)
        line = re.sub(r"^\s*(\d+)\.\s+", r"\1. ", line)
        line = _MARKUP_RE.sub("", line)
        lines.append(line)
    return "\n".join(lines).strip() + "\n"


def _write_docx(path: Path, markdown: str, base_dir: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for raw in (markdown or "").splitlines():
        line = raw.rstrip()
        image = _IMAGE_RE.fullmatch(line.strip())
        if image:
            label = image.group(1) or Path(image.group(2)).name
            image_path = _resolve_image(base_dir, image.group(2))
            if image_path is not None:
                try:
                    document.add_picture(
                        str(image_path),
                        width=Inches(6.0),
                    )
                    if label:
                        paragraph = document.add_paragraph()
                        paragraph.add_run(f"图：{label}").italic = True
                    continue
                except Exception:
                    pass
            document.add_paragraph(f"[图像未嵌入：{label}]")
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            document.add_heading(
                _MARKUP_RE.sub("", heading.group(2)),
                level=min(len(heading.group(1)), 4),
            )
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if bullet:
            document.add_paragraph(
                _MARKUP_RE.sub("", bullet.group(1)),
                style="List Bullet",
            )
            continue
        numbered = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if numbered:
            document.add_paragraph(
                _MARKUP_RE.sub("", numbered.group(2)),
                style="List Number",
            )
            continue
        if line.lstrip().startswith(">"):
            document.add_paragraph(
                _MARKUP_RE.sub("", line.lstrip()[1:].strip())
            )
            continue
        document.add_paragraph(_MARKUP_RE.sub("", line))

    temp = path.with_suffix(".docx.tmp")
    temp.unlink(missing_ok=True)
    document.save(str(temp))
    temp.replace(path)


def _html_document(markdown: str) -> str:
    blocks: list[str] = []
    in_list = False
    for raw in (markdown or "").splitlines():
        line = raw.rstrip()
        if not line:
            if in_list:
                blocks.append("</ul>")
                in_list = False
            blocks.append("<div style='height:6pt'></div>")
            continue
        if _IMAGE_RE.fullmatch(line.strip()):
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            if in_list:
                blocks.append("</ul>")
                in_list = False
            level = min(len(heading.group(1)), 4)
            blocks.append(
                f"<h{level}>{html.escape(heading.group(2))}</h{level}>"
            )
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if bullet:
            if not in_list:
                blocks.append("<ul>")
                in_list = True
            blocks.append(f"<li>{html.escape(bullet.group(1))}</li>")
            continue
        if in_list:
            blocks.append("</ul>")
            in_list = False
        if line.lstrip().startswith(">"):
            value = line.lstrip()[1:].strip()
            blocks.append(
                "<div style='border-left:2px solid #888;padding-left:8px;color:#444'>"
                + html.escape(value)
                + "</div>"
            )
        else:
            blocks.append(
                f"<p>{html.escape(_MARKUP_RE.sub('', line))}</p>"
            )
    if in_list:
        blocks.append("</ul>")
    return (
        "<html><head><meta charset='utf-8'><style>"
        "body{font-family:sans-serif;line-height:1.5;font-size:10.5pt;}"
        "h1{font-size:18pt}h2{font-size:15pt}h3{font-size:12.5pt}"
        "p{margin:0 0 5pt 0}li{margin:0 0 3pt 0}"
        "</style></head><body>"
        + "".join(blocks)
        + "</body></html>"
    )


def _insert_html_checked(page, rect, markup: str, *, scale_low: float) -> None:
    try:
        result = page.insert_htmlbox(
            rect,
            markup,
            scale_low=scale_low,
        )
    except TypeError:
        result = page.insert_htmlbox(rect, markup)
    try:
        spare = float(result[0]) if isinstance(result, tuple) else float(result)
    except Exception:
        return
    if spare < -0.01:
        raise RuntimeError(
            "PDF文本页内容溢出，Phoenix已阻止生成被截断的成品"
        )


def _write_text_page(doc, markdown: str) -> None:
    import fitz

    if not markdown.strip():
        return
    page = doc.new_page(width=595.0, height=842.0)
    rect = fitz.Rect(32.0, 32.0, 563.0, 810.0)
    _insert_html_checked(
        page,
        rect,
        _html_document(markdown),
        scale_low=0.58,
    )


def _write_image_page(doc, image_path: Path, label: str) -> None:
    import fitz

    page = doc.new_page(width=595.0, height=842.0)
    caption_rect = fitz.Rect(32.0, 24.0, 563.0, 54.0)
    caption = (
        "<html><head><meta charset='utf-8'></head>"
        f"<body><b>图：</b>{html.escape(label)}</body></html>"
    )
    _insert_html_checked(
        page,
        caption_rect,
        caption,
        scale_low=0.7,
    )

    target = fitz.Rect(32.0, 60.0, 563.0, 810.0)
    try:
        page.insert_image(
            target,
            filename=str(image_path),
            keep_proportion=True,
        )
    except Exception:
        fallback = (
            "<html><head><meta charset='utf-8'></head>"
            f"<body>[图像无法嵌入：{html.escape(image_path.name)}]</body></html>"
        )
        _insert_html_checked(
            page,
            target,
            fallback,
            scale_low=0.7,
        )


def _pdf_safe_lines(markdown: str):
    for raw in (markdown or "").splitlines():
        if (
            _IMAGE_RE.fullmatch(raw.strip())
            or len(raw) <= _PDF_LINE_CHAR_LIMIT
        ):
            yield raw
            continue
        prefix = ""
        body = raw
        for marker in ("- ", "* ", "+ ", "> "):
            if raw.startswith(marker):
                prefix, body = marker, raw[len(marker):]
                break
        for offset in range(0, len(body), _PDF_LINE_CHAR_LIMIT):
            piece = body[offset : offset + _PDF_LINE_CHAR_LIMIT]
            yield (prefix if offset == 0 else "") + piece


def _atomic_pdf_save(doc, path: Path) -> None:
    path = Path(path)
    temp = path.with_name(path.stem + ".tmp" + path.suffix)
    temp.unlink(missing_ok=True)
    try:
        doc.save(str(temp), garbage=0, deflate=True)
        os.replace(temp, path)
    except Exception:
        temp.unlink(missing_ok=True)
        raise


def _write_pdf(path: Path, markdown: str, base_dir: Path) -> None:
    import fitz

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = fitz.open()
    try:
        current: list[str] = []
        chars = 0

        def flush() -> None:
            nonlocal current, chars
            if current:
                _write_text_page(doc, "\n".join(current))
                current = []
                chars = 0

        for raw in _pdf_safe_lines(markdown):
            image = _IMAGE_RE.fullmatch(raw.strip())
            if image:
                flush()
                label = image.group(1) or Path(image.group(2)).name
                image_path = _resolve_image(base_dir, image.group(2))
                if image_path is not None:
                    _write_image_page(doc, image_path, label)
                else:
                    _write_text_page(doc, f"[图像缺失：{label}]")
                continue

            line_chars = max(len(raw), 1)
            if current and (
                chars + line_chars > _PDF_PAGE_CHAR_LIMIT
                or len(current) >= _PDF_PAGE_LINE_LIMIT
            ):
                flush()
            current.append(raw)
            chars += line_chars
        flush()
        if doc.page_count == 0:
            _write_text_page(doc, "Phoenix医学专题")
        _atomic_pdf_save(doc, path)
    finally:
        doc.close()


class MultiFormatExporter:
    """Export organized medical material to PDF/DOCX/Markdown/TXT with images."""

    def __init__(self, output_root: Path):
        self.output_root = Path(output_root)
        self.output_root.mkdir(parents=True, exist_ok=True)

    def export_text(
        self,
        markdown: str,
        *,
        title: str,
        source_path: Path | None = None,
    ) -> ExportBundle:
        title = _safe_name(title)
        bundle_root = self.output_root / title
        bundle_root.mkdir(parents=True, exist_ok=True)

        md_path = bundle_root / f"{title}.md"
        txt_path = bundle_root / f"{title}.txt"
        docx_path = bundle_root / f"{title}.docx"
        pdf_path = bundle_root / f"{title}.pdf"

        source_base = (
            Path(source_path).parent
            if source_path is not None
            else bundle_root
        )
        md_path.write_text(
            (markdown or "").rstrip() + "\n",
            encoding="utf-8",
        )
        txt_path.write_text(
            markdown_to_plain(markdown),
            encoding="utf-8",
        )
        _write_docx(docx_path, markdown, source_base)
        _write_pdf(pdf_path, markdown, source_base)

        if source_path is not None:
            source_path = Path(source_path)
            source_assets = source_path.with_name(
                source_path.stem + "_assets"
            )
            if source_assets.is_dir():
                target_assets = bundle_root / source_assets.name
                if target_assets.exists():
                    shutil.rmtree(
                        target_assets,
                        ignore_errors=True,
                    )
                shutil.copytree(source_assets, target_assets)

        return ExportBundle(
            output_dir=bundle_root,
            markdown=md_path,
            text=txt_path,
            docx=docx_path,
            pdf=pdf_path,
        )

    def export_path(
        self,
        source: Path,
        *,
        title: str | None = None,
    ) -> ExportBundle:
        source = Path(source)
        if not source.is_file():
            raise FileNotFoundError(source)
        markdown = source.read_text(
            encoding="utf-8",
            errors="replace",
        )
        return self.export_text(
            markdown,
            title=title or source.stem,
            source_path=source,
        )
