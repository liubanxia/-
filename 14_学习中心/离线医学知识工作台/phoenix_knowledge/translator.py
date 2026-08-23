from __future__ import annotations

import html
import json
import os
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from .chunker import chunk_text
from .config import WorkbenchPaths
from .llm import LocalLLM
from .pdf_assets import PDFAssetStore, html_images, markdown_images
from .pdf_parser import iter_pdf_pages, pdf_page_count, sha256_file
from .translation_models import MultiModelTranslationEngine
from .translation_runtime_adapter import TranslationRuntimeAdapter
from .translation_pdf import (
    LAYOUT_ORIGINAL_BILINGUAL,
    LAYOUT_TEXT_BILINGUAL,
    LAYOUT_TRANSLATED_ONLY,
    TranslationPDFBuilder,
)
from .translation_layout_compact import LAYOUT_SOURCE_TRANSLATED


ProgressCallback = Callable[[int, int, str], None]
PauseCallback = Callable[[], bool]
PreviewCallback = Callable[[int, str, Path], None]

EXPORT_PDF = "pdf"
EXPORT_PDF_RICH = "pdf_rich"
EXPORT_RICH = "rich"
EXPORT_TXT = "txt"


def _safe_filename(text: str, max_length: int = 64) -> str:
    text = re.sub(r'[\\/:*?"<>|\r\n]+', '_', text).strip(' ._')
    return (text or '未命名PDF')[:max(8, int(max_length))]


def _normalize_smart_level(value: str | None) -> str:
    raw = (value or "smart1").strip().lower()
    return "smart2" if raw in {"smart2", "2", "deep", "quality", "max"} else "smart1"


def _translation_chunk_chars() -> int:
    raw = os.environ.get("PHOENIX_TRANSLATION_CHUNK_CHARS", "").strip()
    try:
        value = int(raw) if raw else 4800
    except (TypeError, ValueError):
        value = 3200
    return max(2400, min(7200, value))


def _normalize_export_format(value: str | None) -> str:
    raw = (value or EXPORT_PDF).strip().lower()
    if raw in {EXPORT_PDF_RICH, "all", "full"}:
        return EXPORT_PDF_RICH
    if raw in {EXPORT_RICH, "text_rich", "docx"}:
        return EXPORT_RICH
    if raw in {EXPORT_TXT, "text", "plain"}:
        return EXPORT_TXT
    return EXPORT_PDF


def _normalize_layout(value: str | None) -> str:
    if value in {
        LAYOUT_ORIGINAL_BILINGUAL,
        LAYOUT_TEXT_BILINGUAL,
        LAYOUT_TRANSLATED_ONLY,
    }:
        return str(value)
    return LAYOUT_ORIGINAL_BILINGUAL


@dataclass(frozen=True)
class TranslationResult:
    output_path: Path
    source_path: Path
    start_page: int
    total_pages: int
    target_language: str
    pages_done: int
    resumed_pages: int
    warning_pages: int
    available_backends: tuple[str, ...]
    output_paths: tuple[Path, ...] = ()
    image_count: int = 0
    paused: bool = False
    smart_level: str = "smart2"
    output_layout: str = LAYOUT_ORIGINAL_BILINGUAL
    export_format: str = EXPORT_PDF
    part_pages: int = 50


class PDFTranslator:
    """Whole-book offline medical translator.

    Product behavior:
    - formal medical documents always use the Smart2 quality route;
    - every completed page is checkpointed independently;
    - pause occurs safely at a page boundary and resume skips completed pages;
    - the default PDF output preserves images/layout and replaces the source
      text layer in place;
    - one complete PDF is produced by default; split volumes are opt-in.
    """

    def __init__(self, paths: WorkbenchPaths, llm: LocalLLM):
        self.paths = paths
        self.llm = llm
        self.engine = MultiModelTranslationEngine(paths, llm)
        self.translation_runtime = TranslationRuntimeAdapter()
        self.assets = PDFAssetStore(paths.runtime_root)
        self.output_root = paths.evidence_root / 'PDF整本翻译'
        self.output_root.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _write_json(path: Path, payload: dict) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        temp = path.with_suffix(path.suffix + '.tmp')
        temp.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding='utf-8',
        )
        temp.replace(path)

    @staticmethod
    def _read_json(path: Path) -> dict:
        if not path.is_file():
            return {}
        try:
            return json.loads(path.read_text(encoding='utf-8'))
        except Exception:
            return {}

    @staticmethod
    def _resolve_resume_start_page(
        state: dict,
        requested_start_page: int,
        *,
        force_restart: bool = False,
    ) -> int:
        requested = max(1, int(requested_start_page))
        if force_restart or not state:
            return requested
        try:
            existing = int(state.get('start_page', requested))
        except (TypeError, ValueError):
            return requested
        return max(1, existing)

    def _book_paths(
        self,
        pdf_path: Path,
        digest: str,
        target_language: str,
    ) -> tuple[Path, Path, Path, Path, Path]:
        # Preserve existing translation checkpoints when they already exist,
        # but use shorter paths for all new tasks to avoid Windows MAX_PATH
        # failures on long medical textbook titles.
        legacy_name = _safe_filename(pdf_path.stem, 120)
        short_name = _safe_filename(pdf_path.stem, 48)
        language = _safe_filename(target_language, 16)
        legacy_root = self.output_root / f'{legacy_name}_{digest[:10]}_{language}'
        short_root = self.output_root / f'{short_name}_{digest[:10]}_{language}'
        book_root = legacy_root if legacy_root.exists() else short_root

        pages_root = book_root / 'pages'
        audit_root = book_root / 'audit'
        checkpoint = book_root / 'checkpoint.json'
        final_output = book_root / '完整译文.txt'
        pages_root.mkdir(parents=True, exist_ok=True)
        audit_root.mkdir(parents=True, exist_ok=True)
        return book_root, pages_root, audit_root, checkpoint, final_output

    def _translate_page(
        self,
        source_text: str,
        page_number: int,
        target_language: str,
        *,
        smart_level: str = "smart1",
        status: Callable[[str], None] | None = None,
    ) -> tuple[str, dict]:
        source_text = (source_text or '').strip()
        smart_level = _normalize_smart_level(smart_level)
        smart_label = "智能2" if smart_level == "smart2" else "智能1"

        if not source_text:
            # Scanned/image-only pages are kept as their original PDF page by
            # the compact renderer.  They are not a translation-quality
            # failure and must never leak an "人工复核" marker into a delivery.
            return '', {
                'page': page_number,
                'warning_count': 0,
                'skipped_no_text': True,
                'smart_level': smart_level,
                'parts': [{
                    'part': 1,
                    'backend': 'none',
                    'quality_ok': True,
                    'quality_score': 0.0,
                    'reasons': ['PDF未提取到文字；原页已保留'],
                }],
            }

        parts = chunk_text(
            source_text,
            max_chars=_translation_chunk_chars(),
            overlap_chars=0,
        ) or [source_text]
        translated_parts: list[str] = []
        part_audits: list[dict] = []
        warning_count = 0

        for part_index, part_text in enumerate(parts, start=1):
            if status:
                status(
                    f'第 {page_number} 页：{smart_label} 正在精译 '
                    f'{part_index}/{len(parts)} 段'
                )
            try:
                decision = self.translation_runtime.translate(
                    self.engine,
                    part_text,
                    target_language,
                    page_number=page_number,
                    requested_level=smart_level,
                )
                translated = decision.text.strip()
                if decision.needs_review:
                    # The engine already performs a correction pass.  Do not
                    # ship a low-confidence draft or ask the reader to repair
                    # it manually: retain the checkpoint/audit and stop formal
                    # publication below.
                    warning_count += 1
                translated_parts.append(translated)
                part_audits.append({
                    'part': part_index,
                    'part_total': len(parts),
                    'backend': decision.backend,
                    'quality_ok': decision.quality.ok,
                    'quality_score': round(float(decision.quality.score), 4),
                    'reasons': list(decision.quality.reasons),
                    'needs_review': decision.needs_review,
                    'attempts': [
                        {
                            'backend': attempt.backend,
                            'quality_ok': attempt.quality.ok,
                            'quality_score': round(float(attempt.quality.score), 4),
                            'reasons': list(attempt.quality.reasons),
                        }
                        for attempt in decision.attempts
                    ],
                })
                if status:
                    status(
                        f'第 {page_number} 页：第 {part_index}/{len(parts)} 段完成 '
                        f'| 结构校验={decision.quality.score:.2f}'
                    )
            except Exception as exc:
                warning_count += 1
                translated_parts.append('')
                part_audits.append({
                    'part': part_index,
                    'part_total': len(parts),
                    'backend': 'failed_all',
                    'quality_ok': False,
                    'quality_score': 0.0,
                    'reasons': [f'{type(exc).__name__}: {exc}'],
                    'needs_review': True,
                })
                if status:
                    status(
                        f'第 {page_number} 页：第 {part_index}/{len(parts)} 段失败，'
                        '自动重试失败；本次不会发布不合格成品'
                    )

        return '\n\n'.join(translated_parts).strip(), {
            'page': page_number,
            'warning_count': warning_count,
            'smart_level': smart_level,
            'parts': part_audits,
        }

    def _write_preview(
        self,
        pdf_path: Path,
        book_root: Path,
        pages_root: Path,
        start_page: int,
        end_page: int,
        target_language: str,
        *,
        require_all: bool = True,
    ) -> Path:
        path = book_root / '完整译文.txt'
        sections = [
            f'{pdf_path.stem} - {target_language}译文',
            f'原文件：{pdf_path}',
            f'范围：第 {start_page} 页至第 {end_page} 页',
            '',
        ]
        for page_number in range(start_page, end_page + 1):
            page_file = pages_root / f'{page_number:06d}.txt'
            if not page_file.is_file():
                if require_all:
                    raise RuntimeError(
                        f'无法生成完整译文：第 {page_number} 页尚未完成翻译'
                    )
                continue
            translated = page_file.read_text(
                encoding='utf-8', errors='replace'
            ).strip()
            sections.extend([
                f'===== 第 {page_number} 页 =====',
                '',
                translated,
                '',
            ])
        path.write_text('\n'.join(sections).rstrip() + '\n', encoding='utf-8')
        return path

    def _assemble_rich_outputs(
        self,
        pdf_path: Path,
        book_root: Path,
        pages_root: Path,
        start_page: int,
        total_pages: int,
        target_language: str,
    ) -> tuple[tuple[Path, ...], int]:
        txt_path = book_root / '完整译文.txt'
        md_path = book_root / '完整译文.md'
        html_path = book_root / '完整译文.html'
        docx_path = book_root / '完整译文.docx'
        images_root = book_root / 'images'
        images_root.mkdir(parents=True, exist_ok=True)

        txt_sections = [
            f'{pdf_path.stem} - 完整{target_language}译本',
            f'原文件：{pdf_path}',
            f'翻译范围：第 {start_page} 页至第 {total_pages} 页',
            '',
        ]
        md_sections = [
            f'# {pdf_path.stem} - 完整{target_language}译本',
            '',
            f'- 原文件：`{pdf_path}`',
            f'- 翻译范围：第 {start_page} 页至第 {total_pages} 页',
            '',
        ]
        html_sections = [
            '<!doctype html><html><head><meta charset="utf-8">',
            f'<title>{html.escape(pdf_path.stem)} - 完整{html.escape(target_language)}译本</title>',
            '<style>body{font-family:Arial,"Microsoft YaHei",sans-serif;max-width:1000px;margin:2em auto;line-height:1.65;padding:0 1em}pre{white-space:pre-wrap}figure{margin:1.2em 0}img{max-width:100%;border:1px solid #ddd}</style>',
            '</head><body>',
            f'<h1>{html.escape(pdf_path.stem)} - 完整{html.escape(target_language)}译本</h1>',
            f'<p>原文件：{html.escape(str(pdf_path))}</p>',
            f'<p>翻译范围：第 {start_page} 页至第 {total_pages} 页</p>',
        ]

        docx = None
        try:
            from docx import Document
            docx = Document()
            docx.add_heading(f'{pdf_path.stem} - 完整{target_language}译本', 0)
            docx.add_paragraph(f'原文件：{pdf_path}')
            docx.add_paragraph(
                f'翻译范围：第 {start_page} 页至第 {total_pages} 页'
            )
        except Exception:
            docx = None

        image_count = 0
        for page_number in range(start_page, total_pages + 1):
            page_file = pages_root / f'{page_number:06d}.txt'
            if not page_file.is_file():
                raise RuntimeError(
                    f'无法生成整本译本：第 {page_number} 页尚未完成翻译'
                )
            translated = page_file.read_text(encoding='utf-8').strip()
            copied = self.assets.copy_page_assets(
                pdf_path,
                page_number,
                images_root,
                prefix='',
                ensure=True,
            )
            image_count += len(copied)

            txt_sections.extend([
                f'===== 第 {page_number} 页 =====',
                '',
                translated,
                '',
            ])
            md_sections.extend([
                f'## 第 {page_number} 页',
                '',
                translated,
                '',
            ])
            if copied:
                md_sections.append(
                    markdown_images(
                        copied,
                        relative_to=book_root,
                        label=f'第{page_number}页原图',
                    )
                )

            html_sections.extend([
                f'<h2>第 {page_number} 页</h2>',
                f'<pre>{html.escape(translated)}</pre>',
            ])
            if copied:
                html_sections.append(
                    html_images(
                        copied,
                        relative_to=book_root,
                        label=f'第{page_number}页原图',
                    )
                )

            if docx is not None:
                try:
                    docx.add_heading(f'第 {page_number} 页', level=1)
                    docx.add_paragraph(translated)
                    for image_path in copied:
                        try:
                            docx.add_picture(str(image_path))
                        except Exception:
                            docx.add_paragraph(
                                f'[图片文件：{image_path.name}]'
                            )
                except Exception:
                    pass

        html_sections.append('</body></html>')
        txt_path.write_text(
            '\n'.join(txt_sections).rstrip() + '\n',
            encoding='utf-8',
        )
        md_path.write_text(
            '\n'.join(md_sections).rstrip() + '\n',
            encoding='utf-8',
        )
        html_path.write_text('\n'.join(html_sections), encoding='utf-8')
        outputs: list[Path] = [txt_path, md_path, html_path]
        if docx is not None:
            try:
                docx.save(str(docx_path))
                outputs.append(docx_path)
            except Exception:
                pass
        return tuple(outputs), image_count

    # Compatibility surface retained for existing tests / callers.
    def _assemble_outputs(
        self,
        pdf_path: Path,
        book_root: Path,
        pages_root: Path,
        start_page: int,
        total_pages: int,
        target_language: str,
    ) -> tuple[tuple[Path, ...], int]:
        return self._assemble_rich_outputs(
            pdf_path,
            book_root,
            pages_root,
            start_page,
            total_pages,
            target_language,
        )

    def _build_deliverables(
        self,
        pdf_path: Path,
        book_root: Path,
        pages_root: Path,
        start_page: int,
        total_pages: int,
        target_language: str,
        *,
        output_layout: str,
        export_format: str,
        part_pages: int,
        progress: ProgressCallback | None,
    ) -> tuple[tuple[Path, ...], int]:
        outputs: list[Path] = []
        image_count = 0
        export_format = _normalize_export_format(export_format)

        if export_format in {EXPORT_PDF, EXPORT_PDF_RICH}:
            if progress:
                progress(
                    0,
                    total_pages - start_page + 1,
                    '翻译完成，正在生成整书PDF与分册PDF……',
                )
            complete_pdf, part_paths = TranslationPDFBuilder(
                pdf_path,
                pages_root,
                book_root,
            ).build(
                start_page=start_page,
                total_pages=total_pages,
                layout=output_layout,
                part_pages=part_pages,
                progress=progress,
            )
            outputs.append(complete_pdf)
            outputs.extend(part_paths)

        if export_format in {EXPORT_RICH, EXPORT_PDF_RICH}:
            if progress:
                progress(
                    0,
                    total_pages - start_page + 1,
                    '正在生成可编辑文本格式与原图目录……',
                )
            rich_paths, image_count = self._assemble_rich_outputs(
                pdf_path,
                book_root,
                pages_root,
                start_page,
                total_pages,
                target_language,
            )
            outputs.extend(rich_paths)

        if export_format == EXPORT_TXT:
            outputs.append(
                self._write_preview(
                    pdf_path,
                    book_root,
                    pages_root,
                    start_page,
                    total_pages,
                    target_language,
                )
            )

        # Always keep a lightweight local preview/checkpoint text for the GUI,
        # but do not expose it as a selected deliverable unless TXT was chosen.
        self._write_preview(
            pdf_path,
            book_root,
            pages_root,
            start_page,
            total_pages,
            target_language,
        )

        return tuple(dict.fromkeys(outputs)), image_count

    def _paused_result(
        self,
        *,
        pdf_path: Path,
        final_output: Path,
        book_root: Path,
        pages_root: Path,
        start_page: int,
        total_pages: int,
        target_language: str,
        pages_done: int,
        resumed_pages: int,
        warning_pages: int,
        all_backends: list[str],
        smart_level: str,
        output_layout: str,
        export_format: str,
        part_pages: int,
    ) -> TranslationResult:
        self._write_preview(
            pdf_path,
            book_root,
            pages_root,
            start_page,
            total_pages,
            target_language,
            require_all=False,
        )
        return TranslationResult(
            output_path=final_output,
            source_path=pdf_path,
            start_page=start_page,
            total_pages=total_pages,
            target_language=target_language,
            pages_done=pages_done,
            resumed_pages=resumed_pages,
            warning_pages=warning_pages,
            available_backends=tuple(all_backends),
            output_paths=(),
            image_count=0,
            paused=True,
            smart_level=smart_level,
            output_layout=output_layout,
            export_format=export_format,
            part_pages=part_pages,
        )

    def translate_book(
        self,
        pdf_path: Path,
        *,
        start_page: int = 1,
        target_language: str = '中文',
        progress: ProgressCallback | None = None,
        force_restart: bool = False,
        retry_warning_pages: bool = False,
        smart_level: str = "smart2",
        medical_quality_required: bool = True,
        output_layout: str = LAYOUT_ORIGINAL_BILINGUAL,
        export_format: str = EXPORT_PDF,
        part_pages: int = 50,
        should_pause: PauseCallback | None = None,
        page_preview: PreviewCallback | None = None,
    ) -> TranslationResult:
        pdf_path = Path(pdf_path).resolve()
        if not pdf_path.exists() or not pdf_path.is_file():
            raise FileNotFoundError(pdf_path)
        if pdf_path.suffix.lower() != '.pdf':
            raise ValueError(f'仅支持PDF整本翻译: {pdf_path}')

        # Formal PDF translation never permits opting out of medical quality.
        # Keep the legacy argument only for source compatibility.
        del medical_quality_required
        smart_level = "smart2"
        # Public PDF translation is deliberately one-way: PDF in, one full
        # PDF out.  Rich/TXT sidecars and split volumes caused users to open a
        # non-PDF artifact by mistake and also duplicated storage.
        del output_layout, export_format, part_pages
        output_layout = LAYOUT_SOURCE_TRANSLATED
        export_format = EXPORT_PDF
        part_pages = 0

        active_backends = self.engine.active_backends(
            target_language,
            smart_level,
        )
        if not active_backends:
            compute = self.llm.compute.status()
            provider = "本机 Smart2"
            model = self.llm.active_model_name("translation")
            detail = ""
            if self.llm.compute.requested_mode() == "remote":
                try:
                    provider = self.llm.compute.provider_label()
                except Exception:
                    provider = "外接 API"
                detail = (
                    f"当前平台：{provider}；统一模型：{model}。"
                    "请在“模型/算力”重新粘贴该平台 API Key，勾选授权后点击“保存并启用”，"
                    "再点“真实测试算力/API”。"
                )
            raise RuntimeError(
                'Smart2 医学精译未就绪，无法发布正式译文。' + detail
            )

        total_pages = int(pdf_page_count(pdf_path))
        if total_pages <= 0:
            raise RuntimeError('PDF没有可读取页面')

        start_page = max(1, int(start_page))
        if start_page > total_pages:
            raise ValueError(
                f'开始页 {start_page} 超出PDF总页数 {total_pages}'
            )

        digest = sha256_file(pdf_path)
        (
            book_root,
            pages_root,
            audit_root,
            checkpoint_path,
            final_output,
        ) = self._book_paths(pdf_path, digest, target_language)

        if force_restart and book_root.exists():
            for page_file in pages_root.glob('*.txt'):
                page_file.unlink(missing_ok=True)
            for audit_file in audit_root.glob('*.json'):
                audit_file.unlink(missing_ok=True)
            for pattern in ('*.txt', '*.md', '*.html', '*.docx', '*.pdf'):
                for output in book_root.glob(pattern):
                    output.unlink(missing_ok=True)
            shutil.rmtree(book_root / 'images', ignore_errors=True)
            shutil.rmtree(book_root / 'PDF分册', ignore_errors=True)
            checkpoint_path.unlink(missing_ok=True)

        previous_state = self._read_json(checkpoint_path)
        if previous_state:
            if previous_state.get('source_sha256') != digest:
                raise RuntimeError(
                    '检测到PDF内容已变化，请使用重新翻译模式。'
                )
            if previous_state.get('target_language') != target_language:
                raise RuntimeError(
                    '检测到目标语言已变化，请使用重新翻译模式。'
                )

        start_page = self._resolve_resume_start_page(
            previous_state,
            start_page,
            force_restart=force_restart,
        )
        if start_page > total_pages:
            raise RuntimeError(
                f'翻译checkpoint中的开始页 {start_page} 超出PDF总页数 {total_pages}'
            )

        # Old checkpoints were created by the previous raw Marian/NLLB-first
        # policy. Re-translate them once with the new intelligent medical path
        # instead of silently reusing poor text. Checkpoints created by this
        # version resume normally.
        previous_smart = str(previous_state.get('smart_level', '') or '')
        retranslate_existing = bool(
            previous_state
            and not force_restart
            and previous_smart != smart_level
        )

        if progress:
            message = (
                '检测到旧版翻译结果，将按新的智能医学翻译重新精译……'
                if retranslate_existing
                else '正在准备整本医学翻译任务……'
            )
            progress(0, total_pages - start_page + 1, message)

        formal_names = getattr(self.engine, 'formal_backend_names', None)
        all_backends = (
            list(formal_names(target_language))
            if callable(formal_names)
            else [
                str(getattr(backend, 'name', 'medical_translation'))
                for backend in active_backends
            ]
        )
        state = {
            'source_path': str(pdf_path),
            'source_sha256': digest,
            'book_title': pdf_path.stem,
            'target_language': target_language,
            'start_page': start_page,
            'total_pages': total_pages,
            'status': 'running',
            'available_backends': all_backends,
            'smart_level': smart_level,
            'medical_quality_required': True,
            'output_layout': output_layout,
            'export_format': export_format,
            'part_pages': part_pages,
            'last_completed_page': int(
                previous_state.get('last_completed_page', start_page - 1)
            ) if previous_state else start_page - 1,
            'warning_pages': int(
                previous_state.get('warning_pages', 0)
            ) if previous_state else 0,
        }
        self._write_json(checkpoint_path, state)

        selected_total = total_pages - start_page + 1
        pages_done = 0
        resumed_pages = 0
        warning_pages = 0

        try:
            for page_number, source_text in iter_pdf_pages(pdf_path):
                if page_number < start_page:
                    continue

                if should_pause and should_pause():
                    state['status'] = 'paused'
                    state['warning_pages'] = warning_pages
                    self._write_json(checkpoint_path, state)
                    if progress:
                        progress(
                            pages_done,
                            selected_total,
                            '翻译已暂停；已完成页面均已保存，下次可直接继续。',
                        )
                    return self._paused_result(
                        pdf_path=pdf_path,
                        final_output=final_output,
                        book_root=book_root,
                        pages_root=pages_root,
                        start_page=start_page,
                        total_pages=total_pages,
                        target_language=target_language,
                        pages_done=pages_done,
                        resumed_pages=resumed_pages,
                        warning_pages=warning_pages,
                        all_backends=all_backends,
                        smart_level=smart_level,
                        output_layout=output_layout,
                        export_format=export_format,
                        part_pages=part_pages,
                    )

                page_file = pages_root / f'{page_number:06d}.txt'
                audit_file = audit_root / f'{page_number:06d}.json'
                old_audit = self._read_json(audit_file)
                has_warning = int(old_audit.get('warning_count', 0)) > 0
                page_smart = str(old_audit.get('smart_level', '') or '')
                can_resume_page = (
                    page_file.is_file()
                    and page_file.stat().st_size > 0
                    and not retranslate_existing
                    and page_smart == smart_level
                    and not has_warning
                )

                if can_resume_page:
                    resumed_text = page_file.read_text(
                        encoding='utf-8', errors='replace'
                    ).strip()
                    pages_done += 1
                    resumed_pages += 1
                    if has_warning:
                        warning_pages += 1
                    state['last_completed_page'] = max(
                        int(state.get('last_completed_page', start_page - 1)),
                        page_number,
                    )
                    if progress:
                        progress(
                            pages_done,
                            selected_total,
                            f'第 {page_number}/{total_pages} 页已精译，直接续用',
                        )
                    if page_preview:
                        try:
                            page_preview(page_number, resumed_text, page_file)
                        except Exception:
                            pass
                    continue

                def note(message: str) -> None:
                    if progress:
                        progress(pages_done, selected_total, message)

                translated, audit = self._translate_page(
                    source_text,
                    page_number,
                    target_language,
                    smart_level=smart_level,
                    status=note,
                )
                page_file.write_text(
                    translated.rstrip() + '\n',
                    encoding='utf-8',
                )
                self._write_json(audit_file, audit)

                pages_done += 1
                if int(audit.get('warning_count', 0)) > 0:
                    warning_pages += 1
                state['last_completed_page'] = page_number
                state['status'] = 'running'
                state['warning_pages'] = warning_pages
                self._write_json(checkpoint_path, state)
                if page_preview:
                    try:
                        page_preview(page_number, translated, page_file)
                    except Exception:
                        pass
                if progress:
                    progress(
                        pages_done,
                        selected_total,
                        f'整本翻译：已完成第 {page_number}/{total_pages} 页 '
                        f'| 待复核页={warning_pages}',
                    )

                if should_pause and should_pause():
                    state['status'] = 'paused'
                    self._write_json(checkpoint_path, state)
                    if progress:
                        progress(
                            pages_done,
                            selected_total,
                            '已完成当前页并暂停；下次从下一页继续。',
                        )
                    return self._paused_result(
                        pdf_path=pdf_path,
                        final_output=final_output,
                        book_root=book_root,
                        pages_root=pages_root,
                        start_page=start_page,
                        total_pages=total_pages,
                        target_language=target_language,
                        pages_done=pages_done,
                        resumed_pages=resumed_pages,
                        warning_pages=warning_pages,
                        all_backends=all_backends,
                        smart_level=smart_level,
                        output_layout=output_layout,
                        export_format=export_format,
                        part_pages=part_pages,
                    )

            audited_parts = 0
            accepted_parts = 0
            for page_number in range(start_page, total_pages + 1):
                page_audit = self._read_json(
                    audit_root / f'{page_number:06d}.json'
                )
                for part in page_audit.get('parts') or ():
                    if not isinstance(part, dict):
                        continue
                    backend = str(part.get('backend', '') or '')
                    audited_parts += 1
                    if (
                        backend != 'failed_all'
                        and bool(part.get('quality_ok', False))
                    ):
                        accepted_parts += 1
            if audited_parts and accepted_parts != audited_parts:
                raise RuntimeError(
                    '存在未通过医学质量校验的翻译段落；Phoenix已保留逐页'
                    'checkpoint并拒绝发布不合格PDF。请检查模型/API连接后点击继续，'
                    '系统会自动重译失败页，不需要人工修改译文。'
                )

            output_paths, image_count = self._build_deliverables(
                pdf_path,
                book_root,
                pages_root,
                start_page,
                total_pages,
                target_language,
                output_layout=output_layout,
                export_format=export_format,
                part_pages=part_pages,
                progress=progress,
            )
            final_output = output_paths[0] if output_paths else final_output
            state['status'] = 'completed'
            state['last_completed_page'] = total_pages
            state['warning_pages'] = warning_pages
            state['output_path'] = str(final_output)
            state['output_paths'] = [str(x) for x in output_paths]
            state['image_count'] = image_count
            self._write_json(checkpoint_path, state)

            if progress:
                progress(
                    selected_total,
                    selected_total,
                    '整本翻译与PDF分册已完成。',
                )

            return TranslationResult(
                output_path=final_output,
                source_path=pdf_path,
                start_page=start_page,
                total_pages=total_pages,
                target_language=target_language,
                pages_done=pages_done,
                resumed_pages=resumed_pages,
                warning_pages=warning_pages,
                available_backends=tuple(all_backends),
                output_paths=output_paths,
                image_count=image_count,
                paused=False,
                smart_level=smart_level,
                output_layout=output_layout,
                export_format=export_format,
                part_pages=part_pages,
            )
        except Exception as exc:
            state['status'] = 'failed'
            state['error'] = f'{type(exc).__name__}: {exc}'
            self._write_json(checkpoint_path, state)
            raise
        finally:
            self.translation_runtime.clear()
            self.engine.unload()

    def translate(self, pdf_path: Path, **kwargs) -> TranslationResult:
        kwargs.pop('end_page', None)
        return self.translate_book(pdf_path, **kwargs)
